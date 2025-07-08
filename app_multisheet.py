import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from string import ascii_lowercase
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.set_page_config(page_title="Reference List Generator", layout="centered")
st.title("📄 Reference List Generator")

with st.expander("🛈 How to use this tool"):
    st.markdown("[First-time user guide (Word doc)](https://caenergy-my.sharepoint.com/:w:/g/personal/yiming_luo_energy_ca_gov/EZDfTXBgKyxAvGlEmyAf4s4BCiZkErxsJUAIfUAgEyOeCA?e=vgcP3p)")

uploaded_file = st.file_uploader("Upload a spreadsheet (.xlsx or .ods)", type=["xlsx", "ods"])
project_name = st.text_input("File name", "Corby BESS")
prefix = st.text_input("Global Prefix (e.g., CEC, GA, Dudek)", "CEC")
agency_name = st.text_input("Full Agency Name", "California Energy Commission")
proceeding = st.text_input("CEC Proceeding Code (e.g. 24-OPT-05)", "24-OPT-05")
add_header = st.checkbox("Add docket title and headers to top of document")

if add_header:
    user_project_title = st.text_input("Full Project Name, e.g. Corby Battery Energy Storage System Project")
else:
    user_project_title = ""

# Handle workbook + sheet selection
df = None
if uploaded_file:
    try:
        xls = pd.ExcelFile(uploaded_file)
        sheet_names = xls.sheet_names
        selected_sheet = st.selectbox("Select a sheet to use", sheet_names)
        df = pd.read_excel(xls, sheet_name=selected_sheet)
    except Exception as e:
        st.error(f"Could not load workbook: {e}")

def generate_suffixes(n):
    suffixes = []
    reps = 1
    while len(suffixes) < n:
        for ch in ascii_lowercase:
            suffixes.append(ch * reps)
            if len(suffixes) == n:
                break
        reps += 1
    return suffixes

def generate_references(df, prefix, agency_name, proceeding):
    df = df[['TN #', 'Docketed Date', 'Document Title']].dropna()
    df['Year'] = pd.to_datetime(df['Docketed Date'], errors='coerce').dt.year
    df['Formatted Title'] = df['Document Title'].apply(lambda x: str(x).split('\\n')[0].strip())
    df = df.sort_values(by=['Year', 'Docketed Date']).reset_index(drop=True)
    df['GroupIndex'] = df.groupby('Year').cumcount()
    suffix_dict = {y: generate_suffixes(len(g)) for y, g in df.groupby('Year')}
    df['Suffix'] = df.apply(lambda r: suffix_dict[r['Year']][r['GroupIndex']], axis=1)
    df['Formatted Date'] = pd.to_datetime(df['Docketed Date'], errors='coerce').apply(
        lambda x: f"{x.strftime('%B')} {x.day}, {x.year}" if pd.notnull(x) else ""
    )

    base_url = f"https://efiling.energy.ca.gov/Lists/DocketLog.aspx?docketnumber={proceeding}"
    references = [
        (
            f"{prefix} {row['Year']}{row['Suffix']} – {agency_name} (TN {row['TN #']}). {row['Formatted Title']}. Docketed {row['Formatted Date']}. Accessed online at:",
            base_url
        )
        for _, row in df.iterrows()
    ]
    return references

def add_styled_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Tahoma")
    rPr.append(font)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rPr.append(underline)

    run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def build_styled_docx(references, header_lines=None):
    from docx.shared import Inches

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles['Normal']
    normal.font.name = 'Tahoma'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_before = Pt(6)
    normal.paragraph_format.space_after = Pt(6)

    if header_lines:
        for line in header_lines:
            p = doc.add_paragraph(line.upper())
            p.alignment = 1
            run = p.runs[0]
            run.font.name = 'Tahoma'
            run.font.size = Pt(14)
            run.bold = True
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        agency_para = doc.add_paragraph(agency_name)
        run = agency_para.runs[0]
        run.font.name = 'Tahoma'
        run.font.size = Pt(12)
        run.bold = True
        agency_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        agency_para.paragraph_format.space_before = Pt(12)
        agency_para.paragraph_format.space_after = Pt(6)

    for ref_text, url in references:
        p = doc.add_paragraph()
        run = p.add_run(ref_text + " ")
        run.font.name = 'Tahoma'
        run.font.size = Pt(12)
        run.font.underline = False
        run.font.color.rgb = None
        p.paragraph_format.first_line_indent = Pt(-36)
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        add_styled_hyperlink(p, url, url)

    return doc

if df is not None and prefix and agency_name and proceeding:
    try:
        refs = generate_references(df, prefix, agency_name, proceeding)

        header_lines = None
        if add_header and user_project_title.strip():
            header_lines = [
                user_project_title.strip(),
                "DOCKET REFERENCES LIST",
                proceeding
            ]

        doc = build_styled_docx(refs, header_lines=header_lines)
        buffer = BytesIO()
        doc.save(buffer)
        st.success("✅ Reference list ready!")
        st.download_button(
            "📥 Download .docx Reference List",
            buffer.getvalue(),
            file_name=f"{project_name}_References.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"❌ Error: {e}")
      # Placeholder: replace this comment with full app code if needed.
